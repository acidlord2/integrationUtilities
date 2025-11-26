"""
DataMatrix Barcode Generator
Simple DataMatrix barcode generation utility with Excel processing
"""

from PIL import Image, ImageDraw, ImageFont
from pylibdmtx.pylibdmtx import encode
import pandas as pd
import os
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_BREAK
from docxcompose.composer import Composer
from docx2pdf import convert

def generate_datamatrix(data, size=(200, 200), output_path=None, label=None, font_path=None, font_size=8, padding=4):
    """
    Generate a DataMatrix barcode
    
    Args:
        data (str): Data to encode
        size (tuple): Output image size (width, height)
        output_path (str): Path to save the image (optional)
        
    Returns:
        PIL.Image: Generated barcode image
    """
    try:
        # Encode data as DataMatrix
        encoded = encode(data.encode('utf-8'))
        
        # Convert to PIL Image
        img = Image.frombytes('RGB', (encoded.width, encoded.height), encoded.pixels)
        
        # Resize barcode to requested size
        if size != (encoded.width, encoded.height):
            img = img.resize(size, Image.Resampling.NEAREST)

        # If label requested, create a new image with space below
        if label:
            try:
                # Try load custom font or fall back
                if font_path and os.path.exists(font_path):
                    font = ImageFont.truetype(font_path, font_size)
                else:
                    # Attempt common Windows font, fallback to default
                    win_font = r"C:\Windows\Fonts\arial.ttf"
                    if os.path.exists(win_font):
                        font = ImageFont.truetype(win_font, font_size)
                    else:
                        font = ImageFont.load_default()
                # Measure text size with compatibility (prefer textbbox, fallback to textsize)
                measure_draw = ImageDraw.Draw(img)
                if hasattr(measure_draw, 'textbbox'):
                    bbox = measure_draw.textbbox((0,0), label, font=font)
                    text_w, text_h = bbox[2]-bbox[0], bbox[3]-bbox[1]
                else:
                    text_w, text_h = measure_draw.textsize(label, font=font)
                # New image height accommodates text + padding
                new_w = max(img.width, text_w + padding*2)
                new_h = img.height + text_h + padding*3
                combined = Image.new('RGB', (new_w, new_h), color='white')
                # Paste barcode centered horizontally
                barcode_x = (new_w - img.width)//2
                combined.paste(img, (barcode_x, padding))
                draw = ImageDraw.Draw(combined)
                text_x = (new_w - text_w)//2
                text_y = img.height + padding*2
                draw.text((text_x, text_y), label, fill='black', font=font)
                img = combined
            except Exception as e_label:
                print(f"Label render failed: {e_label}")

        if output_path:
            img.save(output_path)
            print(f"DataMatrix saved to: {output_path}")
        return img
        
    except Exception as e:
        print(f"Error generating DataMatrix: {e}")
        return None

def read_xlsx_with_gtin(file_path, target_gtin):
    """
    Read Excel file and extract all rows with matching GTIN
    
    Args:
        file_path (str): Path to Excel file
        target_gtin (str): GTIN to match
        
    Returns:
        list: List of dictionaries containing matched row data
    """
    try:
        # Read Excel file with header in row 1 (index 0) forcing all cells to string to avoid numeric coercion that drops leading zeros
        df = pd.read_excel(file_path, header=0, dtype=str)

        # Print available columns for debugging
        print(f"Available columns: {list(df.columns)}")

        # Find GTIN & Code columns (exact match on normalized header text)
        gtin_column = None
        code_column = None
        for col in df.columns:
            col_lower = str(col).strip().lower()
            if col_lower == 'gtin':
                gtin_column = col
            if col_lower in ('код', 'code', 'qr code'):
                code_column = col
        print (f"Identified GTIN column: {gtin_column}, Code column: {code_column}")
        if gtin_column is None:            
            raise ValueError(f"No GTIN column found and no cells containing '{target_gtin}'")
        if code_column is None:
            raise ValueError(f"No Code column found in the Excel file")

        print(f"Using GTIN column: {gtin_column}; Code column: {code_column}")

        # Normalize GTIN series
        raw_series = df[gtin_column].astype(str).str.strip()
        # Remove internal spaces & non-printable chars
        norm_series = raw_series.str.replace(r'\s+', '', regex=True)
        # Build zero-prefixed variant (if cell length is target_len-1 and target starts with '0')
        target_str = str(target_gtin).strip()
        target_no0 = target_str.lstrip('0')
        target_len = len(target_str)
        # Create a series with potential leading zero restored
        restored_series = norm_series.apply(lambda v: ('0'+v) if len(v)==target_len-1 and target_str.startswith('0') else v)

        # Build match mask (exact, restored, no leading zero variant, or contains)
        match_mask = (
            norm_series.eq(target_str) |
            restored_series.eq(target_str) |
            norm_series.eq(target_no0) |
            restored_series.eq(target_no0) |
            norm_series.str.contains(target_no0, na=False)
        )

        matched_df = df[match_mask]

        if matched_df.empty:
            # Prepare debug samples
            sample = norm_series.head(10).tolist()
            restored_sample = restored_series.head(10).tolist()
            print("GTIN debug samples (raw -> restored):")
            for a,b in zip(sample, restored_sample):
                print(f"  {a} -> {b}")
            raise ValueError(
                f"No rows matched GTIN '{target_gtin}'. Tried variants: '{target_str}', '{target_no0}'. "
                f"First 10 normalized values: {sample}"
            )

        matched_rows = []
        for idx, (_, row) in enumerate(matched_df.iterrows(), start=1):
            gtin_val = str(row[gtin_column]).strip()
            # Restore leading zero if needed for reporting
            if target_str.startswith('0') and len(gtin_val)==target_len-1:
                gtin_val = '0'+gtin_val
            code_val = str(row[code_column]).strip() if code_column in row else gtin_val
            filename = f"dm_{target_str}_{idx}.png"
            matched_rows.append({'gtin': gtin_val, 'code': code_val, 'filename': filename})
        
        print(f"Found {len(matched_rows)} matching rows")
        return matched_rows
        
    except Exception as e:
        print(f"Error reading Excel file: {e}")
        return []

def generate_barcodes_for_rows(matched_rows, gtin, output_dir="barcodes", size=(50,50), overwrite=False):
    """Generate DataMatrix barcodes for each matched row.

    Files will be saved under output_dir named: dm_<gtin>_<index>.png

    Args:
        matched_rows (list): List of dicts each having 'code' key.
        gtin (str): GTIN used for naming.
        output_dir (str): Directory to store barcodes.
        size (tuple): Barcode image size.
    Returns:
        list: Paths of generated barcode image files.
    """
    os.makedirs(output_dir, exist_ok=True)
    safe_gtin = str(gtin).strip()
    generated = []
    for idx, row in enumerate(matched_rows, start=1):
        code = row.get('code') or row.get('gtin')
        if not code:
            print(f"Row {idx} missing code, skipping")
            continue
        filename = row.get('filename') or f"dm_{safe_gtin}_{idx}.png"
        path = os.path.join(output_dir, filename)
        if not overwrite and os.path.exists(path):
            generated.append(path)
            continue
        img = generate_datamatrix(code, size, path, label=code, font_size=8)
        if img:
            generated.append(path)
    print(f"Generated {len(generated)} barcode files in '{output_dir}'")
    return generated

def append_codes_to_template(matched_rows, barcodes_dir="barcodes", result_dir="result"):
    """Create a single DOCX document with all barcodes.

    Args:
        matched_rows (list): List of dicts with 'code' and 'filename'.
        barcodes_dir (str): Directory where barcode PNGs are stored.
        result_dir (str): Output directory for final DOCX.
    Returns:
        str: Path to the generated DOCX file or None on failure.
    """
    try:
        if not matched_rows:
            print("No matched rows provided; skipping DOCX generation")
            return None
        os.makedirs(result_dir, exist_ok=True)

        print(f"Creating DOCX with {len(matched_rows)} barcodes")

        # Create new document
        doc = Document()
        p = doc.add_paragraph()
        
        # Add all barcode images
        for row in matched_rows:
            filename = row.get('filename')
            image_path = os.path.join(barcodes_dir, filename) if filename else None
            if image_path and os.path.exists(image_path):
                run_img = p.add_run()
                run_img.add_picture(image_path, width=Cm(3))
            else:
                p.add_run("[Image not found]\n")
                        
        output_filename = f"all_barcodes.docx"
        output_path = os.path.join(result_dir, output_filename)
        doc.save(output_path)
        print(f"Saved DOCX: {output_path}")
        return output_path
    except Exception as e:
        print(f"Error generating DOCX: {e}")
        return None

if __name__ == "__main__":
    # Demo with Excel processing
    input_file = "QR codes Miss Korea.xlsx"
    gtin_array = ["04560119224880", "08801114415078", "08802580872266"]
    all_rows = []
    print(f"Processing Excel file: {input_file}")
    for idx, gtin in enumerate(gtin_array):
        print(f"Looking for GTIN: {gtin}")
    
        if os.path.exists(input_file):
            matched_rows = read_xlsx_with_gtin(input_file, gtin)
            if matched_rows:
                all_rows.extend(matched_rows)
                print(f"\nFound {len(matched_rows)} matching rows:")
                for i, row in enumerate(matched_rows, 1):
                    print(f"Row {i}: GTIN={row['gtin']}, Code={row['code']}")
                # Generate barcodes (only once for all matched rows)
                generate_barcodes_for_rows(matched_rows, gtin, output_dir="barcodes", size=(50,50))
            else:
                print("No matching rows found")
        else:
            print(f"File {input_file} not found")

    # Create batch DOCX files then merge them into one DOCX
    if all_rows:
        doc_file = append_codes_to_template(all_rows, barcodes_dir="barcodes", result_dir="result")
        convert(doc_file, os.path.join("result", f"output_{gtin_array}.pdf"))
    else:
        print("No matched rows found in total; skipping DOCX generation.")
        
